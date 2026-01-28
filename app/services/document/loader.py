import re
from typing import Union
from pathlib import Path

import logging

logger = logging.getLogger(__name__)

class DocumentLoader:
    """
    Service for loading and cleaning text from various file types (PDF, TXT, DOCX).
    """

    def load(self, file_path: Union[str, Path]) -> list:
        from app.schemas.document import Document
        
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._load_pdf(path)
        elif suffix == ".txt":
            return self._load_txt(path)
        elif suffix in [".docx", ".doc"]:
            return self._load_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _load_pdf(self, path: Path) -> list:
        documents = []
        from app.schemas.document import Document
        
        # METHOD 1: Try PyMuPDF (Fast)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            logger.info(f"Using PyMuPDF to load {path.name}")
            
            for i, page in enumerate(doc):
                try:
                    text = page.get_text()
                    if text:
                        cleaned_page_text = self._clean_text(text, page_num=i+1)
                        if cleaned_page_text:
                            documents.append(Document(
                                text=cleaned_page_text,
                                metadata={"page": i+1, "source": path.name}
                            ))
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {i+1} of {path.name}: {e}")
                    continue
            doc.close()
            return documents
            
        except ImportError:
            logger.warning("PyMuPDF (fitz) not found. Falling back to pypdf (slower).")
        except Exception as e:
            logger.error(f"PyMuPDF failed: {e}. Falling back to pypdf.")

        # METHOD 2: Fallback to pypdf
        try:
            from pypdf import PdfReader
            logger.info(f"Using pypdf to load {path.name}")
            reader = PdfReader(str(path))
            
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        cleaned_page_text = self._clean_text(text, page_num=i+1)
                        if cleaned_page_text:
                            documents.append(Document(
                                text=cleaned_page_text,
                                metadata={"page": i+1, "source": path.name}
                            ))
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {i+1} of {path.name}: {e}")
                    continue
            return documents
            
        except Exception as e:
            logger.error(f"Error reading PDF {path} with pypdf: {e}")
            raise ValueError(f"Failed to read PDF: {e}. Ensure 'pymupdf' or 'pypdf' is installed.")

    def _load_txt(self, path: Path) -> list:
        from app.schemas.document import Document
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
            
            if not text:
                return []
                
            cleaned_text = self._clean_text(text)
            if not cleaned_text:
                return []

            # TXT files don't have pages, so we treat it as Page 1 or split?
            # For simplicity, treating as single document for now, chunker will handle it.
            return [Document(
                text=cleaned_text,
                metadata={"page": 1, "source": path.name}
            )]
        except Exception as e:
            logger.error(f"Error reading TXT {path}: {e}")
            raise ValueError(f"Failed to read TXT file: {e}")

    def _load_docx(self, path: Path) -> list:
        from app.schemas.document import Document
        try:
            import docx  # python-docx
            doc = docx.Document(path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            text = "\n".join(full_text)
            if not text:
                return []

            cleaned_text = self._clean_text(text)
            if not cleaned_text:
                return []
                
            return [Document(
                text=cleaned_text, 
                metadata={"page": 1, "source": path.name}
            )]
        except ImportError:
            raise ImportError("python-docx not installed. Run `pip install python-docx`")
        except Exception as e:
            logger.error(f"Error reading DOCX {path}: {e}")
            raise ValueError(f"Failed to read DOCX file: {e}")

    def _clean_text(self, text: str, page_num: int = 0) -> str:
        """
        Apply heuristics to clean extracted text.
        """
        # 1. Normalize whitespace
        text = re.sub(r'[ \t]+', ' ', text)
        
        # 2. Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if self._is_page_number(line, page_num):
                continue
                
            if len(line) < 4 and not line[0].isalnum():
                 continue

            cleaned_lines.append(line)
            
        return "\n".join(cleaned_lines)

    def _is_page_number(self, line: str, page_num: int) -> bool:
        if line.isdigit() and (int(line) == page_num or int(line) < 1000):
            return True
        if re.match(r'^(page|pg)\.?\s*\d+$', line, re.IGNORECASE):
            return True
        if re.match(r'^-\s*\d+\s*-$', line):
            return True
        return False
